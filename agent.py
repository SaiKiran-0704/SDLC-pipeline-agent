from dotenv import load_dotenv
from google import genai
from google.genai import types
import json
import time
from docx import Document

load_dotenv()

client = genai.Client()  # reads GEMINI_API_KEY from env automatically

BRD_SCHEMA = {
    "type": "object",
    "properties": {
        "title": {"type": "string"},
        "executive_summary": {"type": "string"},
        "business_objectives": {
            "type": "array",
            "items": {"type": "string"}
        },
        "in_scope": {
            "type": "array",
            "items": {"type": "string"}
        },
        "out_of_scope": {
            "type": "array",
            "items": {"type": "string"}
        },
        "functional_requirements": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "id": {"type": "string"},
                    "description": {"type": "string"},
                    "priority": {"type": "string", "enum": ["Must", "Should", "Could"]}
                },
                "required": ["id", "description", "priority"]
            }
        },
        "assumptions": {
            "type": "array",
            "items": {"type": "string"}
        },
        "open_questions": {
            "type": "array",
            "items": {"type": "string"}
        }
    },
    "required": ["title", "executive_summary", "business_objectives", "in_scope", "out_of_scope", "functional_requirements", "assumptions", "open_questions"]
}

SYSTEM_PROMPT = """You are a Business Requirements Analyst.
Given a raw, possibly messy business request, produce a clear draft Business Requirements Document covering:
- Executive Summary
- Business Objectives
- Scope (In Scope / Out of Scope)
- Functional Requirements
- Assumptions
Do not invent facts you weren't given. If something is unclear, note it as an open question instead of guessing."""

MAX_RETRIES = 2
REQUEST_TIMEOUT_MS = 30_000  # 30 seconds


def generate_brd_json(user_request: str, codebase_context: dict | None = None) -> dict:
    """Takes a raw business request, returns a validated BRD as a dict.
    Retries on both bad JSON and failed/hung requests."""
    contents = user_request
    if codebase_context:
        tree_summary = "\n".join(codebase_context["file_tree"][:50])
        contents += (
            f"\n\nThis feature is being added to an EXISTING codebase. "
            f"Here is its file structure:\n{tree_summary}\n\n"
            f"Ground your requirements in this reality — reference existing "
            f"patterns where relevant, and don't assume things need to be "
            f"built from scratch if they likely already exist."
        )

    last_error = None
    for attempt in range(1, MAX_RETRIES + 2):
        try:
            response = client.models.generate_content(
                model="gemini-3.5-flash",
                contents=contents,
                config=types.GenerateContentConfig(
                    system_instruction=SYSTEM_PROMPT,
                    response_mime_type="application/json",
                    response_schema=BRD_SCHEMA,
                    http_options=types.HttpOptions(timeout=REQUEST_TIMEOUT_MS)
                )
            )
            if not response.text:
                raise ValueError("Model returned an empty response")
            return json.loads(response.text)

        except (json.JSONDecodeError, ValueError) as e:
            print(f"[attempt {attempt}] invalid JSON: {e}")
            last_error = e
            contents = (
                f"{user_request}\n\n"
                f"Your previous response was not valid JSON. Error: {e}\n"
                f"Return ONLY valid JSON matching the schema, nothing else."
            )

        except Exception as e:
            print(f"[attempt {attempt}] request failed: {e}")
            last_error = e
            time.sleep(1)  # brief pause before retrying a network-level failure

    raise RuntimeError(f"Agent failed after {MAX_RETRIES + 1} attempts: {last_error}")


def render_docx(data: dict, output_path: str):
    """Takes a BRD dict, writes it to a formatted .docx file."""
    doc = Document()

    doc.add_heading(data["title"], level=0)

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(data["executive_summary"])

    doc.add_heading("Business Objectives", level=1)
    for obj in data["business_objectives"]:
        doc.add_paragraph(obj, style="List Bullet")

    doc.add_heading("Scope", level=1)
    doc.add_heading("In Scope", level=2)
    for item in data["in_scope"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Out of Scope", level=2)
    for item in data["out_of_scope"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Functional Requirements", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"

    header_cells = table.rows[0].cells
    header_cells[0].text = "ID"
    header_cells[1].text = "Description"
    header_cells[2].text = "Priority"

    for fr in data["functional_requirements"]:
        row_cells = table.add_row().cells
        row_cells[0].text = fr["id"]
        row_cells[1].text = fr["description"]
        row_cells[2].text = fr["priority"]

    doc.add_heading("Assumptions", level=1)
    for a in data["assumptions"]:
        doc.add_paragraph(a, style="List Bullet")

    doc.add_heading("Open Questions", level=1)
    for q in data["open_questions"]:
        doc.add_paragraph(q, style="List Bullet")

    doc.save(output_path)
    print(f"Saved {output_path}")


if __name__ == "__main__":
    user_request = input("Describe your business requirement: ").strip()
    if not user_request:
        print("No request entered. Exiting.")
        exit(1)

    data = generate_brd_json(user_request)
    render_docx(data, "BRD.docx")
    print(json.dumps(data, indent=2))