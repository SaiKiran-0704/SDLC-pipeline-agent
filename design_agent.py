from dotenv import load_dotenv
from google import genai
from google.genai import types
import json
from docx import Document

load_dotenv()
client = genai.Client()

DESIGN_SCHEMA = {
    "type": "object",
    "properties": {
        "title": {"type": "string"},
        "architecture_overview": {"type": "string"},
        "components": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "name": {"type": "string"},
                    "responsibility": {"type": "string"},
                    "addresses_requirement_ids": {
                        "type": "array",
                        "items": {"type": "string"}
                    }
                },
                "required": ["name", "responsibility", "addresses_requirement_ids"]
            }
        },
        "data_model": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "entity": {"type": "string"},
                    "key_fields": {"type": "array", "items": {"type": "string"}}
                },
                "required": ["entity", "key_fields"]
            }
        },
        "tech_stack_suggestions": {"type": "array", "items": {"type": "string"}},
        "design_risks": {"type": "array", "items": {"type": "string"}},
        "open_questions": {"type": "array", "items": {"type": "string"}}
    },
    "required": ["title", "architecture_overview", "components", "data_model", "tech_stack_suggestions", "design_risks", "open_questions"]
}

DESIGN_SYSTEM_PROMPT = """You are a Solutions Architect. You will be given a complete
Business Requirements Document as JSON — not a raw request, a fully structured BRD
with functional requirements, scope, and priorities already defined.

Your job: produce a technical design that satisfies it.

Rules:
1. Every functional requirement marked "Must" in the BRD must be addressed by at
   least one component. Use the exact requirement IDs (e.g. "FR-1") in each
   component's addresses_requirement_ids field — this is how traceability works,
   don't skip it or leave it vague.
2. Don't invent technology choices ungrounded in the request. If the BRD didn't
   specify a platform, database, or existing system, suggest reasonable options
   in tech_stack_suggestions but flag the ambiguity in open_questions instead of
   silently assuming.
3. design_risks are technical risks specifically — performance, security,
   integration complexity, scalability — not business risks (those already
   live in the BRD).
4. Keep architecture_overview to 3-5 sentences: the shape of the solution, not
   an exhaustive spec.
"""

MAX_RETRIES = 2

def generate_design_json(brd_data: dict, feedback: str | None = None, codebase_context: dict | None = None) -> dict:
    """Takes a validated BRD dict, returns a validated Design dict."""
    brd_as_text = json.dumps(brd_data, indent=2)
    contents = f"Here is the BRD:\n\n{brd_as_text}"
    if feedback:
        contents += f"\n\nRevision feedback from reviewer: {feedback}\nPlease regenerate the design, incorporating this feedback."
    if codebase_context:
        tree_summary = "\n".join(codebase_context["file_tree"][:50])
        contents += (
            f"\n\nThis feature is being added to an EXISTING codebase. "
            f"Here is its file structure:\n{tree_summary}\n\n"
            f"Design components that fit this existing structure — reuse "
            f"existing patterns, naming conventions implied by the file "
            f"tree, and don't propose a parallel architecture that ignores "
            f"what's already there."
        )

    for attempt in range(1, MAX_RETRIES + 2):
        try:
            response = client.models.generate_content(
                model="gemini-3.5-flash",
                contents=contents,
                config=types.GenerateContentConfig(
                    system_instruction=DESIGN_SYSTEM_PROMPT,
                    response_mime_type="application/json",
                    response_schema=DESIGN_SCHEMA,
                    http_options=types.HttpOptions(timeout=60_000)
                )
            )
            if not response.text:
                raise ValueError("Model returned an empty response")
            return json.loads(response.text)
        except (json.JSONDecodeError, ValueError) as e:
            print(f"[attempt {attempt}] invalid JSON: {e}")
            last_error = e
            contents += f"\n\nYour previous response was invalid JSON: {e}\nReturn ONLY valid JSON."
        except Exception as e:
            print(f"[attempt {attempt}] request failed: {e}")
            last_error = e

    raise RuntimeError(f"Design agent failed after {MAX_RETRIES + 1} attempts: {last_error}")


def render_design_docx(data: dict, output_path: str):
    """Takes a Design dict, writes it to a formatted .docx file."""
    doc = Document()

    doc.add_heading(data["title"], level=0)

    doc.add_heading("Architecture Overview", level=1)
    doc.add_paragraph(data["architecture_overview"])

    doc.add_heading("Components", level=1)
    for c in data["components"]:
        doc.add_heading(c["name"], level=2)
        doc.add_paragraph(c["responsibility"])
        ids = ", ".join(c["addresses_requirement_ids"])
        doc.add_paragraph(f"Addresses: {ids}", style="Intense Quote")

    doc.add_heading("Data Model", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Entity"
    header_cells[1].text = "Key Fields"
    for entity in data["data_model"]:
        row_cells = table.add_row().cells
        row_cells[0].text = entity["entity"]
        row_cells[1].text = ", ".join(entity["key_fields"])

    doc.add_heading("Suggested Tech Stack", level=1)
    for t in data["tech_stack_suggestions"]:
        doc.add_paragraph(t, style="List Bullet")

    doc.add_heading("Design Risks", level=1)
    for r in data["design_risks"]:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_heading("Open Questions", level=1)
    for q in data["open_questions"]:
        doc.add_paragraph(q, style="List Bullet")

    doc.save(output_path)
    print(f"Saved {output_path}")


if __name__ == "__main__":
    from agent import generate_brd_json

    user_request = input("Describe your business requirement: ").strip()
    if not user_request:
        print("No request entered. Exiting.")
        exit(1)

    print("\n[Requirements Agent] generating BRD...\n")
    brd = generate_brd_json(user_request)
    print(f"BRD created: \"{brd['title']}\" ({len(brd['functional_requirements'])} functional requirements)\n")
    brd_fr_ids = {fr["id"] for fr in brd["functional_requirements"]}
    print(f"BRD requirement IDs: {sorted(brd_fr_ids)}\n")

    print("[Design Agent] generating design from BRD...\n")
    design = generate_design_json(brd)

    design_fr_ids = set()
    for c in design["components"]:
        design_fr_ids.update(c["addresses_requirement_ids"])

    missing = brd_fr_ids - design_fr_ids
    unknown = design_fr_ids - brd_fr_ids

    if missing:
        print(f"⚠️  BRD requirements NOT addressed by any component: {sorted(missing)}")
    if unknown:
        print(f"⚠️  Design references requirement IDs that don't exist in the BRD: {sorted(unknown)}")
    if not missing and not unknown:
        print("✅ Traceability check passed — every BRD requirement is covered, no invented IDs.")

    render_design_docx(design, "Design.docx")
    print(json.dumps(design, indent=2))